"""Word出力モジュール（python-docx）"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_paragraph(doc: Document, text: str):
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())


def _insert_visuals(doc: Document, visuals: list[dict]):
    """章に紐づくビジュアルをWordに挿入する"""
    for v in visuals:
        png_path = v.get("path")
        caption = v.get("caption", "")
        ai_prompt = v.get("ai_prompt", "")

        if png_path and Path(png_path).exists():
            # PNG画像を挿入
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(png_path), width=Cm(14))
            except Exception:
                pass
            if caption:
                cap_p = doc.add_paragraph(caption)
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # PNG生成できなかった場合：ビジュアル枠だけ置く
            doc.add_paragraph(f"【図: {v.get('title', '')}】（画像を差し込んでください）")
            if caption:
                doc.add_paragraph(caption)

        # AIプロンプトは下部に参考として記載
        if ai_prompt:
            note = doc.add_paragraph(f"＜画像AI生成プロンプト＞\n{ai_prompt}")
            note.runs[0].font.size = Pt(8)
            note.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _add_footer(doc: Document, copyright_text: str):
    """全セクションのフッターにコピーライトを設定する"""
    if not copyright_text:
        return
    for section in doc.sections:
        footer = section.footer
        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()
        p.clear()
        run = p.add_run(copyright_text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def export_report(
    output_path: Path,
    candidate: dict,
    cover_assets: dict,
    greeting: str,
    intro: str,
    sections: list[dict],         # [{"title": str, "text": str, "number": int}]
    conclusion: str,
    terms: str,
    section_visuals: dict = None, # {section_number: [{"path": Path, "caption": str, ...}]}
    author_name: str = "",
    copyright_text: str = "",
    profile_page: str = "",
):
    """
    レポート全体をWord形式で書き出す。
    section_visuals: 各章のビジュアルリスト（任意）
    """
    doc = Document()

    # ---- フッター設定（全ページ共通）----
    if copyright_text:
        _add_footer(doc, copyright_text)

    # ---- 表紙 ----
    title_text = cover_assets.get("title", candidate.get("title", ""))
    subtitle_text = cover_assets.get("subtitle", candidate.get("subtitle", ""))
    badges = cover_assets.get("badges", candidate.get("badges", []))

    # タイトル（大きく・中央・太字・ネイビー）
    doc.add_paragraph()  # 上部の余白
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title_text)
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    doc.add_paragraph()

    # サブタイトル
    if subtitle_text:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(subtitle_text)
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    # バッジ
    if badges:
        badge_text = "　".join(f"【{b}】" for b in badges)
        bp = doc.add_paragraph(badge_text)
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bp.runs:
            bp.runs[0].font.size = Pt(11)
            bp.runs[0].font.bold = True
            bp.runs[0].font.color.rgb = RGBColor(0xe8, 0xa0, 0x20)

    doc.add_paragraph()
    doc.add_paragraph()

    # 著者名
    if author_name:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(author_name)
        author_run.font.size = Pt(12)
        author_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    # ---- 表紙デザイン指示書（構造化） ----
    design_lines = ["＜表紙デザイン指示書＞", ""]

    design_lines.append("タイトル:")
    design_lines.append(f"- {title_text}")
    design_lines.append("")

    if subtitle_text:
        design_lines.append("サブタイトル:")
        design_lines.append(f"- {subtitle_text}")
        design_lines.append("")

    if badges:
        design_lines.append("バッジ:")
        for b in badges:
            design_lines.append(f"- 【{b}】")
        design_lines.append("")

    design_lines.append("デザイン:")
    color_scheme = cover_assets.get("color_scheme", "")
    if color_scheme:
        design_lines.append(f"- カラー：{color_scheme}")
    subject = cover_assets.get("subject", "")
    if subject:
        design_lines.append(f"- 被写体：{subject}")
    design_lines.append("- サイズ：A4縦")
    design_lines.append("")

    design_lines.append("表紙文字サイズと比率:")
    design_lines.append("- 1. メインタイトル：最大サイズ（60-80pt相当）。紙面の3分の1を占める程度のインパクト。中央または上部に配置。")
    design_lines.append("- 2. サブキャッチコピー：中サイズ（30-40pt相当）。タイトルの補足として、メインの半分程度の大きさで配置。")
    design_lines.append("- 3. 著者名・補足情報：最小サイズ（18-24pt相当）。可読性を保ちつつ、四隅や下部に控えめに配置。")
    design_lines.append("")

    design_concept = cover_assets.get("design_concept", "")
    if design_concept:
        design_lines.append("デザインコンセプト:")
        design_lines.append(design_concept)
        design_lines.append("")

    image_prompt = cover_assets.get("image_prompt", "")
    if image_prompt:
        design_lines.append("表紙画像生成プロンプト（Midjourney / DALL-E / Firefly 等に貼り付けてご使用ください）:")
        design_lines.append(image_prompt)

    note = doc.add_paragraph("\n".join(design_lines))
    if note.runs:
        note.runs[0].font.size = Pt(8)
        note.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ---- あいさつ ----
    _add_heading(doc, "はじめに（著者より）", level=1)
    _add_paragraph(doc, greeting)
    doc.add_page_break()

    # ---- はじめに ----
    _add_heading(doc, "はじめに", level=1)
    _add_paragraph(doc, intro)
    doc.add_page_break()

    # ---- 本文各章 ----
    for section in sections:
        num = section["number"]
        visuals = (section_visuals or {}).get(num, [])

        # 章開始前ビジュアル
        before = [v for v in visuals if v.get("position") == "before_text"]
        if before:
            _insert_visuals(doc, before)

        _add_heading(doc, f"第{num}章　{section['title']}", level=1)
        _add_paragraph(doc, section["text"])

        # 章末ビジュアル
        after = [v for v in visuals if v.get("position") != "before_text"]
        if after:
            _insert_visuals(doc, after)

        doc.add_page_break()

    # ---- おわりに ----
    _add_heading(doc, "おわりに", level=1)
    _add_paragraph(doc, conclusion)
    doc.add_page_break()

    # ---- 著者プロフィール（入力がある場合のみ）----
    if profile_page:
        _add_heading(doc, "著者プロフィール", level=1)
        _add_paragraph(doc, profile_page)
        doc.add_page_break()

    # ---- 利用規約 ----
    _add_heading(doc, "特典利用規約", level=1)
    _add_paragraph(doc, terms)

    doc.save(str(output_path))
    return output_path
